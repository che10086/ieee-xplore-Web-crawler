# 作者CheJiaQi github:https://github.com/che10086
# 版本1.0
# 运行ieee爬虫需要安装requests库和python-docx库
# python运行版本推荐3.8以上
# 实现功能：输入想要搜索的ieee关键词，输出搜索到的所有文章的标题，摘要，摘要翻译，检索日期，关键词，文章主页，输出文件为word。
# 在运行程序前，请申请一个有道翻译私人API，填入trans.py文件开头的应用ID和应用密钥.

from get_articlelist import airtcle_list
from get_information import single_information, get_title, get_abstract, get_published, get_date, get_kwd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from trans import connect

# 输入搜索关键词
searchword = input('Search:')

# 获得文章列表
Airtcle_List = airtcle_list(searchword)
list_long = len(Airtcle_List)

# 获取每篇文章信息并保存到word中
document = Document()
document.styles['Normal'].font.name = u'等线'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')
# word输入开头信息
document.add_heading('搜索内容：' + searchword, 0)
p = document.add_paragraph()
run = p.add_run("共搜索到" + str(list_long) + '篇文章')
run.bold = False
run.font.name = '等线'
run.font.size = Pt(16)

n = 1  # 记录第几篇文章
for airtcle_number in Airtcle_List:
    page_text = single_information(airtcle_number)
    title = get_title(page_text)
    print("\r", '读取文章:' + title + ' ' + str(n) + '/' + str(list_long), end="", flush=True)
    n = n + 1  # 用来计数

    # 写入title
    p = document.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    # 写入abstract
    p = document.add_paragraph()
    run = p.add_run('Abstract:')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(get_abstract(page_text))
    run.bold = False

    # 写入翻译
    p = document.add_paragraph()
    run = p.add_run('翻译:')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(connect(get_abstract(page_text)))
    run.bold = False

    # 写入出版处
    p = document.add_paragraph()
    run = p.add_run('Published in:')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(get_published(page_text))
    run.bold = False

    # 写入检索日期
    p = document.add_paragraph()
    run = p.add_run('Date Added to IEEE Xplore:')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(get_date(page_text))
    run.bold = False

    # 写入文章关键词
    p = document.add_paragraph()
    run = p.add_run('Keywords:')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(get_kwd(page_text))
    run.bold = False

    # 写入文章地址
    p = document.add_paragraph()
    run = p.add_run('Airtcle URL:')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('https://ieeexplore.ieee.org/document/' + str(airtcle_number))
    run.bold = False

    document.add_page_break()

document.save(searchword.replace(" ", "_") + '.docx')
print('\n', 'Finish!')
